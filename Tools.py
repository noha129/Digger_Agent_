# tools.py
import json
import requests
from io import BytesIO
from bs4 import BeautifulSoup
import PyPDF2
import docx

from langchain_classic.agents import Tool
from langchain_community.utilities import DuckDuckGoSearchAPIWrapper
# -------- Tools --------
ddg_wrapper = DuckDuckGoSearchAPIWrapper()

def duckduckgo_search(query: str) -> str:
    """Search DuckDuckGo and return formatted results."""
    print(f"[DuckDuckGo] Searching: {query}")
    try:
        results = ddg_wrapper.results(query, max_results=5)
        formatted = [
            {
                "title": r.get("title", "N/A"),
                "url": r.get("link", "N/A"),
                "snippet": r.get("snippet", "")
            }
            for r in results
        ]
        return json.dumps(formatted, indent=2)
    except Exception as e:
        return f"Search failed: {str(e)}"


def youtube_search_tool(query: str) -> dict:
    """Search YouTube and return top 4 results."""
    print(f"[YouTube] Searching: {query}")
    try:
        yt = YouTubeSearchTool()
        raw_results = yt.run(query)
        results_data = json.loads(raw_results)

        videos = [
            {
                "title": r.get("title"),
                "url": r.get("link"),
                "snippet": r.get("snippet", "")
            }
            for r in results_data
        ]

        return {
            "results": {
                "youtube_videos": videos
            }
        }
    except Exception as e:
        return {
            "results": {
                "youtube_videos": [],
                "error": str(e)
            }
        }


def fetch_url(url: str) -> str:
    """Fetch webpage content."""
    print(f"[Fetch URL] Fetching: {url}")
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, 'html.parser')
        for tag in soup(['script', 'style']):
            tag.decompose()

        text = soup.get_text(separator='\n', strip=True)
        return text[:5000] + ("\n[Truncated...]" if len(text) > 5000 else "")
    except Exception as e:
        return f"Error: {str(e)}"


def extract_pdf_text(file_path: str) -> str:
    """Extract text from a PDF file (local path or URL)."""
    print(f"[PDF Extractor] Processing: {file_path}")
    try:
        if file_path.startswith('http://') or file_path.startswith('https://'):
            response = requests.get(file_path, timeout=30)
            response.raise_for_status()
            pdf_file = BytesIO(response.content)
        else:
            pdf_file = open(file_path, 'rb')

        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_content = []

        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text.strip():
                text_content.append(f"--- Page {page_num + 1} ---\n{text}")

        full_text = "\n\n".join(text_content)

        if not file_path.startswith('http'):
            pdf_file.close()

        max_chars = 15000
        if len(full_text) > max_chars:
            return full_text[:max_chars] + f"\n\n[Truncated... Total pages: {len(pdf_reader.pages)}]"

        return full_text
    except Exception as e:
        return f"Error extracting PDF: {str(e)}"


def extract_docx_text(file_path: str) -> str:
    """Extract text from a DOCX file (local path or URL)."""
    print(f"[DOCX Extractor] Processing: {file_path}")
    try:
        if file_path.startswith('http://') or file_path.startswith('https://'):
            response = requests.get(file_path, timeout=30)
            response.raise_for_status()
            docx_file = BytesIO(response.content)
        else:
            docx_file = file_path

        doc = docx.Document(docx_file)
        text_content = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        full_text = "\n\n".join(text_content)

        max_chars = 15000
        if len(full_text) > max_chars:
            return full_text[:max_chars] + f"\n\n[Truncated... Total paragraphs: {len(doc.paragraphs)}]"

        return full_text
    except Exception as e:
        return f"Error extracting DOCX: {str(e)}"

